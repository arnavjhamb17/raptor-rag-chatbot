import fitz
import os
import docx

def chunk_pdf(pdf_path):
    chunks = []
    doc = fitz.open(pdf_path)
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text()
        if text.strip():
            chunks.append({
                "text": text,
                "metadata": {
                    "source_file": os.path.basename(pdf_path),
                    "chunk_index": page_num + 1,
                    "chunk_type": "pdf_page"
                }
            })
    return chunks

def chunk_docx(docx_path):
    doc = docx.Document(docx_path)
    chunks = []
    table_rows = []

    # === Chunk paragraphs ===
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            chunks.append({
                "text": text,
                "metadata": {
                    "source_file": os.path.basename(docx_path),
                    "chunk_index": i,
                    "chunk_type": "paragraph"
                }
            })

    # === Collect all table rows ===
    for table_idx, table in enumerate(doc.tables):
        rows = table.rows
        headers = [cell.text.strip() for cell in rows[0].cells] if rows else []
        for row_idx, row in enumerate(rows[1:], start=1):  # skip header row
            cells = [cell.text.strip() for cell in row.cells]
            row_str = " | ".join(cells)
            if row_str.strip():
                table_rows.append(row_str)

    # === Group all table rows into a single table_cluster chunk ===
    if table_rows:
        table_text = "\n".join(table_rows)
        chunks.append({
            "text": table_text,
            "metadata": {
                "source_file": os.path.basename(docx_path),
                "chunk_type": "table_cluster",
                "extraction_method": "docx_table_cluster",
                "is_legal_text": True
            }
        })

    return chunks

# Remove script execution and saving to chunks.json.
# Let unified_rag_pipeline.py handle orchestration and saving.